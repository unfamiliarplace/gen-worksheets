from __future__ import annotations
from docx import Document
from docx.text.paragraph import Paragraph
from typing import Iterable, Type
import docx2pdf
from pathlib import Path
from docx.table import _Cell

class Worksheet:

    @staticmethod
    def prompt_options() -> dict:
        return {}

    @staticmethod
    def parse_data() -> dict:
        return {}
    
    @staticmethod
    def reset() -> None:
        return
    
    @staticmethod
    def make(tag: str='', data: dict={}, opts: dict={}) -> Document:
        raise NotImplementedError

    @classmethod
    def test(cls: Type[Worksheet], path_output: Path) -> None:        
        opts = cls.prompt_options()
        data = cls.parse_data()
        cls.reset()
        d = cls.make('test', data, opts)
        path_output.mkdir(parents=True, exist_ok=True)
        d.save(path_output / '_test.docx')
        # docx2pdf.convert(path_output / '_test.docx', path_output / '_test.pdf')
        print(f'Saved test file in {path_output}')

    @staticmethod
    def fill_cell(c: _Cell, val: str) -> None:
        c.paragraphs[0].text = val

    @staticmethod
    def replace(d: Document, key: str, val: str, limit: int=0) -> int:
        """
        Replace the given placeholder with the given value.
        Stop after finding limit instances. Supply 0 (default) for unlimited.
        Return the number of instances found and replaced.
        """
        key_ = f'__{key}__'
        found = 0
        
        def _yield_paras() -> Iterable[Paragraph]:
            for s in d.sections:
                for p in s.header.paragraphs:
                    yield p
                for p in s.footer.paragraphs:
                    yield p
            for p in d.paragraphs:
                yield p

        for p in _yield_paras():
            if p.text.find(key_) >= 0:
                p.text = p.text.replace(key_, val)  
                found += 1
                if (limit > 0) and (found >= limit):
                    break
        
        return found
