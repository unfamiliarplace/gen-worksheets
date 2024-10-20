import random
import math
from worksheet import Worksheet
from docx import Document
from docx.table import Table
from pathlib import Path

path_template = Path('templates/binary.docx')
path_output = Path('output/binary')
path_test = path_output / 'test.docx'

PH_TAG = '__TAG__'

# Dumb and potentially infinite... TODO
used_default = {
    'dec_to_bin': set(),
    'bin_to_dec': set(),
    'dec_to_power': set(),
    'power_to_states': set(),
}

used = {
    'dec_to_bin': set(),
    'bin_to_dec': set(),
    'dec_to_power': set(),
    'power_to_states': set(),
}

def dec_to_bin(lower: int, upper: int) -> tuple[str]:
    n = random.randint(lower, upper)
    while n in used['dec_to_bin']:
        n = random.randint(lower, upper)
    used['dec_to_bin'].add(n)

    return str(n), str(bin(n))[2:]

def bin_to_dec(lower: int, upper: int) -> tuple[str]:
    answer, n = dec_to_bin(lower, upper)
    while n in used['bin_to_dec']:
        answer, n = dec_to_bin(lower, upper)
    used['bin_to_dec'].add(n)

    return n, answer

def dec_to_power(lower: int, upper: int) -> tuple[str]:
    n = random.randint(lower, upper)
    while n in used['dec_to_power']:
        n = random.randint(lower, upper)
    used['dec_to_power'].add(n)

    return str(n), str(math.ceil(math.log2(n)))

def power_to_states(lower: int, upper: int) -> tuple[str]:
    n = random.randint(lower, upper)
    while n in used['power_to_states']:
        n = random.randint(lower, upper)
    used['power_to_states'].add(n)

    return str(n), str(n ** 2)

def fill_table(t: Table, coords: tuple[int], func: callable, args: list=[], kwargs: dict={}) -> None:
    for coord in coords:
        cell = t.cell(*coord)
        cell.text = func(*args, **kwargs)[0]

class Binary(Worksheet):    

    @staticmethod
    def reset() -> None:
        for k in used:
            used[k] = set()

    @staticmethod
    def make(tag: str='', data: dict={}) -> None:

        d = Document(path_template)
        tables = d.tables
        coords = ((0, 0), (1, 0), (0, 2), (1, 2))

        fill_table(tables[0], coords, bin_to_dec, [0, 64])
        fill_table(tables[1], coords, dec_to_bin, [0, 64])
        fill_table(tables[2], coords, dec_to_power, [1, 8])
        fill_table(tables[3], coords, power_to_states, [1, 256])
        
        # tag it
        for p in d.paragraphs:
            if p.text.find(PH_TAG) >= 0:
                p.text = p.text.replace(PH_TAG, tag)

        return d

if __name__ == '__main__':
    Binary.test(path_output)
